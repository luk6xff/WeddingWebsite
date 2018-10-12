#!/usr/bin/python3

# install python-docx: https://python-docx.readthedocs.io/en/latest/user/install.html
# $ pip install python-docx


from docx import Document
from docx.shared import Mm
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_TAB_ALIGNMENT
import wedding_gsheet_downloader as wgd


def create_name_table_tags_docx_for_guests_who_confirmed():
    all_guests = wgd.generate_names_dict(wgd.extract_guests_from_wedding_sheet())
    if all_guests is None:
        raise Exception('guests_dict is empty!')    
    # extract people who confirmed attending at wedding party
    confirmed_list = []
    for i, _ in enumerate(all_guests ['name']):
        if all_guests ['invited'][i] == 'T' and all_guests ['confirmed'][i] == 'T':
            confirmed_list.append({'name':all_guests ['name'][i],
                                   'surname':all_guests ['surname'][i]}) 
    


    # create doc
    doc = Document()
    cols_num = 2
    table = doc.add_table(rows=0, cols=cols_num)
    table.alignment = WD_TAB_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i in range(0, len(confirmed_list), cols_num):
        row = table.add_row()
        # set cell width and height
        for cell in row.cells:
            cell.width  = Mm(80)
            cell.height = Mm(55)
        for col_num in range(cols_num):
            prgs = row.cells[col_num].paragraphs
            img_width = Mm(70)
            img_height = Mm(15)
            prgs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            prgs[0].add_run().add_picture('flowers_name_tag_up.png', height=img_height, width=img_width)
            text_run = prgs[0].add_run()
            text_run.font.size = Pt(14)
            text_run.font.name = 'Cambria'
            text_run.bold = True
            text_run.italic = True
            try:
                text_run.add_text('{} {}'.format(confirmed_list[i]['name'], confirmed_list[i]['surname']))
            except:
                pass
            prgs[0].add_run().add_picture('flowers_name_tag_down.png', height=img_height, width=img_width)
            i = i+1

    doc.add_page_break()
    doc.save("table_name_tags.docx")


if __name__ == "__main__":
    create_name_table_tags_docx_for_guests_who_confirmed()
